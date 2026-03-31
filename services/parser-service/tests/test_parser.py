"""Unit tests for parser-service worker (text extraction + section structuring)."""
import io
import pytest

from worker import extract_text_docx, structure_sections


# ── Text extraction ───────────────────────────────────────────────────────────

def test_structure_sections_basic():
    sample_text = """
John Doe

Skills
Python, FastAPI, Docker, Redis

Experience
Software Engineer at Acme Corp 2021-2024
Built microservices using Python and FastAPI.

Education
B.Sc. Computer Science, MIT, 2021
"""
    sections = structure_sections(sample_text)
    assert "skills" in sections
    assert "experience" in sections
    assert "education" in sections
    assert "Python" in sections["skills"]
    assert "Acme Corp" in sections["experience"]


def test_structure_sections_empty():
    sections = structure_sections("")
    assert sections == {}


def test_structure_sections_no_headers():
    text = "This is a plain text resume with no recognisable section headers."
    sections = structure_sections(text)
    # Falls into 'other'
    assert "other" in sections


# ── DOCX extraction ───────────────────────────────────────────────────────────

def test_extract_text_docx_basic():
    """Build a minimal DOCX in-memory and verify text extraction."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Skills: Python, FastAPI")
    doc.add_paragraph("Experience: 3 years")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    text = extract_text_docx(buf.read())
    assert "John Doe" in text
    assert "Python" in text
    assert "Experience" in text
